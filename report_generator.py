from __future__ import annotations

import os
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd

from database import fetch_reports, fetch_tasks, fetch_validation_logs, fetch_villages
from schema import CORE_REQUIRED_FIELDS, FIELD_LABELS, INDICATOR_FIELDS


MISSING_TABLE_COLUMNS = [
    "Thôn",
    "Kỳ báo cáo",
    "Nhóm dữ liệu",
    "Trường cần bổ sung/sửa",
    "Tình trạng hiện tại",
    "Mức độ",
    "Gợi ý xử lý",
    "Nguồn",
]

METADATA_LABELS = {
    "commune_name": "Tên xã",
    "village_name": "Tên thôn",
    "period": "Kỳ báo cáo",
    "reporter_name": "Người lập báo cáo",
    "reporter_title": "Chức danh người lập",
    "phone": "Số điện thoại",
    "due_at": "Hạn nộp",
    "submitted_at": "Thời điểm nộp",
    "submission_status": "Trạng thái nộp",
}


def _normalize(value: Any) -> str:
    """Normalize Vietnamese text for lightweight rule matching."""
    text = "" if value is None else str(value)
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return "_".join(text.lower().strip().split())


def _is_missing(value: Any) -> bool:
    if value is None:
        return True
    try:
        if pd.isna(value):
            return True
    except (TypeError, ValueError):
        pass
    return str(value).strip() == ""


def _field_label(field_name: str) -> str:
    if field_name in FIELD_LABELS:
        return f"{field_name[:4].upper()} - {FIELD_LABELS[field_name]}"
    return METADATA_LABELS.get(field_name, field_name or "Không xác định")


def _field_group(field_name: str) -> str:
    if field_name in INDICATOR_FIELDS:
        return "Chỉ tiêu CT01-CT14"
    if field_name in {"submission_status", "due_at", "submitted_at"}:
        return "Tiến độ nộp báo cáo"
    return "Thông tin báo cáo"


def _suggestion(field_name: str, severity: str, error_type: str = "") -> str:
    error_norm = _normalize(error_type)
    if field_name == "submission_status" or "chua_nop" in error_norm:
        return "Nộp báo cáo đúng kỳ và cập nhật trạng thái trên hệ thống."
    if field_name == "phone":
        return "Cập nhật số điện thoại gồm 10 chữ số và bắt đầu bằng số 0."
    if field_name in INDICATOR_FIELDS:
        return f"Bổ sung hoặc kiểm tra lại giá trị {_field_label(field_name)} trong file nguồn."
    if severity == "BLOCKER":
        return "Bổ sung/sửa dữ liệu bắt buộc rồi kiểm định lại trước khi lưu."
    return "Rà soát thông tin và cập nhật lại khi có số liệu chính xác."


def build_summary_dataframe(period: str) -> pd.DataFrame:
    """Build one dashboard row per village with all CT01-CT14 fields."""
    villages = pd.DataFrame(fetch_villages())
    reports = pd.DataFrame(fetch_reports(period))
    tasks = pd.DataFrame(fetch_tasks(period))

    if villages.empty:
        return pd.DataFrame()

    out = villages[["commune_name", "village_name"]].copy()

    requested_report_cols = [
        "village_name",
        "period",
        "phone",
        "reporter_name",
        "reporter_title",
        "due_at",
        "submitted_at",
        "submission_status",
        "days_late",
        *INDICATOR_FIELDS,
    ]

    if not reports.empty:
        available_report_cols = [
            column for column in requested_report_cols if column in reports.columns
        ]
        report_view = reports[available_report_cols].copy()
        # Protect the merge if a legacy database contains duplicate rows.
        report_view = report_view.drop_duplicates(subset=["village_name"], keep="last")
        out = out.merge(report_view, on="village_name", how="left")
    else:
        out["period"] = period

    task_cols = [
        "village_name",
        "status",
        "due_at",
        "submitted_at",
        "days_late",
        "reminder_status",
    ]
    if not tasks.empty:
        available_task_cols = [column for column in task_cols if column in tasks.columns]
        task_view = tasks[available_task_cols].copy()
        task_view = task_view.drop_duplicates(subset=["village_name"], keep="last")

        # Avoid suffix ambiguity and prefer task timing/status as the operational source.
        out = out.merge(task_view, on="village_name", how="left", suffixes=("", "_task"))
        for field in ["due_at", "submitted_at", "days_late"]:
            task_field = f"{field}_task"
            if task_field in out.columns:
                if field not in out.columns:
                    out[field] = out[task_field]
                else:
                    out[field] = out[task_field].combine_first(out[field])
                out = out.drop(columns=[task_field])

    if "period" not in out.columns:
        out["period"] = period
    else:
        out["period"] = out["period"].fillna(period)

    if "status" in out.columns:
        task_status = out["status"]
        if "submission_status" in out.columns:
            out["submission_status"] = task_status.combine_first(out["submission_status"])
        else:
            out["submission_status"] = task_status
        out = out.drop(columns=["status"])
    elif "submission_status" not in out.columns:
        out["submission_status"] = "Chưa nộp"

    out["submission_status"] = out["submission_status"].fillna("Chưa nộp")

    # Guarantee all indicators exist so the dashboard can render CT01-CT14
    # consistently even when an older database has no value for a field yet.
    for field in INDICATOR_FIELDS:
        if field not in out.columns:
            out[field] = pd.NA

    return out.sort_values("village_name", kind="stable").reset_index(drop=True)


def build_missing_fields_dataframe(period: str) -> pd.DataFrame:
    """Return actionable missing/invalid fields for the selected period."""
    summary = build_summary_dataframe(period)
    logs = pd.DataFrame(fetch_validation_logs(period))
    rows: list[dict[str, Any]] = []

    # Derive missing submissions and missing core CT01-CT04 values directly
    # from the operational data, even when validation logs were cleared.
    if not summary.empty:
        for _, record in summary.iterrows():
            village = str(record.get("village_name") or "Không xác định")
            status = str(record.get("submission_status") or "Chưa nộp")
            status_norm = _normalize(status)

            if "chua_nop" in status_norm:
                rows.append(
                    {
                        "Thôn": village,
                        "Kỳ báo cáo": period,
                        "Nhóm dữ liệu": "Tiến độ nộp báo cáo",
                        "Trường cần bổ sung/sửa": "Báo cáo kỳ hiện tại",
                        "Tình trạng hiện tại": "Chưa nộp",
                        "Mức độ": "BLOCKER",
                        "Gợi ý xử lý": "Nộp file báo cáo của kỳ hiện tại để hệ thống kiểm định và tổng hợp.",
                        "Nguồn": "Suy luận từ trạng thái nhiệm vụ",
                    }
                )
                # A not-submitted village naturally has no CT values yet; avoid
                # creating 4 extra duplicate rows for the same root cause.
                continue

            for field in CORE_REQUIRED_FIELDS:
                if field in {"commune_name", "village_name", "period"}:
                    continue
                if field in record.index and _is_missing(record.get(field)):
                    rows.append(
                        {
                            "Thôn": village,
                            "Kỳ báo cáo": period,
                            "Nhóm dữ liệu": _field_group(field),
                            "Trường cần bổ sung/sửa": _field_label(field),
                            "Tình trạng hiện tại": "Đang để trống",
                            "Mức độ": "BLOCKER",
                            "Gợi ý xử lý": _suggestion(field, "BLOCKER"),
                            "Nguồn": "Suy luận từ dữ liệu tổng hợp",
                        }
                    )

    # Add relevant AI Validator logs: missing fields, not-submitted records,
    # and invalid formats. Other statistical warnings remain in the quality log.
    if not logs.empty:
        summary_by_village = (
            summary.set_index("village_name", drop=False) if not summary.empty else None
        )
        for _, log in logs.iterrows():
            error_type = str(log.get("error_type") or "")
            message = str(log.get("message") or "")
            combined_norm = _normalize(f"{error_type} {message}")
            if not any(
                token in combined_norm
                for token in ["thieu", "chua_nop", "sai_dinh_dang", "khong_hop_le"]
            ):
                continue

            village = str(log.get("village_name") or "Không xác định")
            field_name = str(log.get("field_name") or "")
            severity = str(log.get("severity") or "WARNING").upper()
            current_value = "Cần kiểm tra"

            if summary_by_village is not None and village in summary_by_village.index:
                record = summary_by_village.loc[village]
                if isinstance(record, pd.DataFrame):
                    record = record.iloc[-1]
                if field_name and field_name in record.index:
                    value = record.get(field_name)
                    current_value = "Đang để trống" if _is_missing(value) else str(value)
                elif field_name == "submission_status":
                    current_value = str(record.get("submission_status") or "Chưa nộp")

            rows.append(
                {
                    "Thôn": village,
                    "Kỳ báo cáo": str(log.get("period") or period),
                    "Nhóm dữ liệu": _field_group(field_name),
                    "Trường cần bổ sung/sửa": _field_label(field_name),
                    "Tình trạng hiện tại": current_value,
                    "Mức độ": severity if severity in {"BLOCKER", "WARNING"} else "WARNING",
                    "Gợi ý xử lý": _suggestion(field_name, severity, error_type),
                    "Nguồn": str(log.get("source_file") or "AI Validator"),
                }
            )

    if not rows:
        return pd.DataFrame(columns=MISSING_TABLE_COLUMNS)

    result = pd.DataFrame(rows, columns=MISSING_TABLE_COLUMNS)
    result = result.drop_duplicates(
        subset=["Thôn", "Kỳ báo cáo", "Nhóm dữ liệu", "Trường cần bổ sung/sửa"],
        keep="last",
    )
    severity_order = pd.Categorical(
        result["Mức độ"], categories=["BLOCKER", "WARNING"], ordered=True
    )
    result = (
        result.assign(_severity_order=severity_order)
        .sort_values(["_severity_order", "Thôn", "Trường cần bổ sung/sửa"], kind="stable")
        .drop(columns=["_severity_order"])
        .reset_index(drop=True)
    )
    return result


def generate_excel_report(period: str, output_path: str | Path) -> str:
    output_path = str(output_path)
    parent = os.path.dirname(output_path)
    if parent:
        os.makedirs(parent, exist_ok=True)

    reports = pd.DataFrame(fetch_reports(period))
    tasks = pd.DataFrame(fetch_tasks(period))
    logs = pd.DataFrame(fetch_validation_logs(period))
    summary = build_summary_dataframe(period)
    missing = build_missing_fields_dataframe(period)

    if not reports.empty:
        rename_map = {
            field: f"{field[:4].upper()} - {FIELD_LABELS[field]}"
            for field in INDICATOR_FIELDS
        }
        reports_export = reports.rename(columns=rename_map)
    else:
        reports_export = reports

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        summary.to_excel(writer, sheet_name="Tong_quan", index=False)
        reports_export.to_excel(writer, sheet_name="So_lieu_chi_tiet", index=False)
        tasks.to_excel(writer, sheet_name="Tien_do", index=False)
        missing.to_excel(writer, sheet_name="Du_lieu_thieu", index=False)
        logs.to_excel(writer, sheet_name="Canh_bao", index=False)

        for worksheet in writer.sheets.values():
            worksheet.freeze_panes = "A2"
            for cell in worksheet[1]:
                cell.font = cell.font.copy(bold=True)
            for column in worksheet.columns:
                max_len = max(len(str(cell.value or "")) for cell in column)
                worksheet.column_dimensions[column[0].column_letter].width = min(
                    max(max_len + 2, 12), 40
                )
    return output_path


def _add_table(doc, df: pd.DataFrame, columns: list[str], labels: list[str]) -> None:
    if df.empty:
        doc.add_paragraph("Không có dữ liệu.")
        return
    available = [column for column in columns if column in df.columns]
    if not available:
        doc.add_paragraph("Không có dữ liệu phù hợp.")
        return
    available_labels = [labels[columns.index(column)] for column in available]
    table = doc.add_table(rows=1, cols=len(available))
    table.style = "Table Grid"
    for index, label in enumerate(available_labels):
        table.rows[0].cells[index].text = label
    for _, row in df[available].fillna("").iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(available):
            cells[index].text = str(row[column])


def generate_word_report(period: str, output_path: str | Path) -> str:
    output_path = str(output_path)
    parent = os.path.dirname(output_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("Thiếu python-docx. Chạy: pip install python-docx") from exc

    reports = pd.DataFrame(fetch_reports(period))
    tasks = pd.DataFrame(fetch_tasks(period))
    logs = pd.DataFrame(fetch_validation_logs(period))
    summary = build_summary_dataframe(period)
    missing = build_missing_fields_dataframe(period)

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    doc.add_heading("BÁO CÁO TỔNG HỢP VĂN HÓA - XÃ HỘI THEO THÔN", level=1)
    doc.add_paragraph(f"Kỳ báo cáo: {period}")
    doc.add_paragraph(f"Ngày tạo báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    total_villages = len(summary)
    submitted = (
        int((summary.get("submission_status", pd.Series(dtype=str)) != "Chưa nộp").sum())
        if not summary.empty
        else 0
    )
    total_households = (
        int(reports.get("ct01_households", pd.Series(dtype=float)).sum())
        if not reports.empty
        else 0
    )
    total_population = (
        int(reports.get("ct02_population", pd.Series(dtype=float)).sum())
        if not reports.empty
        else 0
    )
    total_poor = (
        int(reports.get("ct03_poor_households", pd.Series(dtype=float)).sum())
        if not reports.empty
        else 0
    )

    doc.add_heading("1. Tổng quan", level=2)
    for item in [
        f"Tổng số thôn: {total_villages}",
        f"Số thôn đã nộp: {submitted}",
        f"Số thôn chưa nộp: {total_villages - submitted}",
        f"Tổng số hộ dân đã ghi nhận: {total_households}",
        f"Tổng số nhân khẩu đã ghi nhận: {total_population}",
        f"Tổng số hộ nghèo đã ghi nhận: {total_poor}",
    ]:
        doc.add_paragraph(item)

    doc.add_heading("2. Chi tiết theo thôn", level=2)
    detail_columns = ["village_name", *INDICATOR_FIELDS, "submission_status"]
    detail_labels = [
        "Thôn",
        *[f"{field[:4].upper()} - {FIELD_LABELS[field]}" for field in INDICATOR_FIELDS],
        "Trạng thái",
    ]
    _add_table(doc, reports, detail_columns, detail_labels)

    doc.add_heading("3. Tiến độ nộp báo cáo", level=2)
    _add_table(
        doc,
        tasks,
        ["village_name", "due_at", "submitted_at", "status", "days_late", "reminder_status"],
        ["Thôn", "Hạn nộp", "Thời điểm nộp", "Trạng thái", "Số ngày trễ", "Nhắc việc"],
    )

    doc.add_heading("4. Dữ liệu thiếu cần bổ sung", level=2)
    _add_table(doc, missing, MISSING_TABLE_COLUMNS, MISSING_TABLE_COLUMNS)

    doc.add_heading("5. Cảnh báo chất lượng dữ liệu", level=2)
    _add_table(
        doc,
        logs,
        ["village_name", "severity", "error_type", "field_name", "message"],
        ["Thôn", "Mức độ", "Loại lỗi", "Trường", "Nội dung"],
    )

    doc.add_paragraph("Báo cáo được tạo tự động bởi Ba Na SmartLink MVP.")
    doc.save(output_path)
    return output_path
